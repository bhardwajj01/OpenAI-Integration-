from docx import Document
from openai import OpenAI
import time
import os

# Set your OpenAI API key here (replace 'your_openai_api_key' with your actual key)
os.environ['OPENAI_API_KEY'] = 'your_openai_api_key'

chat_history = []
client = OpenAI()

def chat(chat_history):
    # Call OpenAI's chat completion API
    response = client.chat.completions.create(
        model="gpt-4-1106-preview",
        messages=chat_history,
    )
    return response.choices[0].message.content.strip()

def extract_bold_sections(doc_path):
    # Function to extract bold sections from a Word document
    # ...

# Add the Path of your input DOCX file
doc_path = '/path/to/your/input/docx/file'
bold_sections = extract_bold_sections(doc_path)

output_prefix = 'gpt_responses'
output_format = '.docx'

output_directory = 'output_folder'
if not os.path.exists(output_directory):
    os.makedirs(output_directory)

# User prompt to be used in the conversation with the GPT model
prompt = (
    "Below is educational text about accounting fraud. "
    "I would like you to convert below to a lengthy mystery novel genre based in Canada without any chapter breaks. "
    "Each text will have mostly dialogue and continue the story seamlessly. Make it mostly dialogue. "
    "Be detailed and don't skip any key information. Make the text sound like a dialogue-heavy mystery novel. "
    "The audience are accountants and auditors. The novel must educate the readers. "
    "The result should be like a final book that a reader reads without any script information, "
    "and without any division into chapters. In the second and future prompts, continue the story smoothly without restarting. "
    "Don't explain what you write, simply give the output."
)

prev_user_prompt = None

for index, (bold_text, paragraph_text) in enumerate(bold_sections.items(), start=1):
    new_doc = Document()
    output_filename = f"{output_prefix}_{index}{output_format}"
    output_path = os.path.join(output_directory, output_filename)
    chat_prompt = f"{prompt}\n\n{bold_text}\n{paragraph_text}\n"

    # Displaying the current chat prompt
    print("======================Chatprompt========================")
    print(chat_prompt)
   
    # Building chat history for the conversation with GPT model
    if prev_user_prompt:
        chat_history = [prev_user_prompt]
        chat_history.append({"role": "assistant", "content": prev_gpt_response})
        chat_history.append({"role": "user", "content": chat_prompt})
    else:
        chat_history.append({"role": "user", "content": chat_prompt})
    
    
    # Get GPT response for the current prompt
    gpt_response = chat(chat_history)
    prev_gpt_response = gpt_response
    prev_user_prompt = {"role": "user", "content": chat_prompt}
    
    # Save the GPT response to the new Word document
    new_doc.add_paragraph(f"{gpt_response}")
    new_doc.add_paragraph("\n")
    new_doc.save(output_path)
    
    # Introducing a delay to avoid rate-limiting issues
    time.sleep(5)
